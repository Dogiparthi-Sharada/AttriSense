# ---------------------------------------------------------------------------
# AttriSense — scripts/build_beginners_docx.py
# ---------------------------------------------------------------------------
# Author : Sharada Dogiparthi <dogiparthi.sharada@gmail.com>
# Version: 1.0.0
# Date   : 2026-05-07
# License: MIT — see LICENSE in repo root.
# Copyright (c) 2026 Sharada Dogiparthi. All rights reserved.
# ---------------------------------------------------------------------------
"""Build a Beginner's Guide DOCX from docs/learn/ markdown sources.

Polished output:
  - Cover page with title, tagline, version, date
  - Auto-generated Word TOC (updates on open in Word)
  - Per-chapter title page with key idea + accent rule
  - Inline rendering of headings, bullets, ordered lists, code, blockquotes,
    bold/italic/inline-code, links, math blocks, and pipe tables
  - Embedded Pixel-pastel diagram PNGs
  - Captioned images, clean spacing, page numbers in footer
  - Google Sans / Inter (falls back to Calibri / Consolas if missing)

Run:
    .venv/bin/python scripts/build_beginners_docx.py
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
DOCS = REPO / "docs"
OUTPUTS = REPO / "outputs"
OUTPUTS.mkdir(parents=True, exist_ok=True)
OUT = OUTPUTS / "AttriSense_Beginners_Guide.docx"

# ---------------- Pixel pastel ink colours
INK = RGBColor(0x1F, 0x1E, 0x1B)
BODY = RGBColor(0x3D, 0x3A, 0x35)
MUTED = RGBColor(0x7A, 0x74, 0x6B)
LAVENDER = RGBColor(0x6E, 0x58, 0xB8)
PEACH = RGBColor(0xC9, 0x7A, 0x4A)
SAGE = RGBColor(0x5C, 0x8B, 0x4F)
CODE_INK = RGBColor(0x5A, 0x4E, 0x3F)

ACCENT_HEX = "B8B5E1"        # lavender
TINT_HEX = "F5EFE6"          # sand
CARD_HEX = "FBF7F2"          # cream
CODE_BG_HEX = "F5EFE6"       # sand for code blocks

SECTIONS = [
    # md, ascii companion, diagram png, chapter title, key-idea
    ("learn/index.md", None, None, "Welcome",
     "A friendly tour of every technique used in the AttriSense dashboard."),
    ("learn/shap.md", "ascii/shap.txt", "diagrams/data-flow.png",
     "SHAP — Explain Any Prediction",
     "Decompose one prediction into the contribution of each feature."),
    ("learn/smote.md", "ascii/smote.txt", "diagrams/ml-pipeline.png",
     "SMOTE — Class Balance",
     "Synthesize minority-class examples — but only on the train fold."),
    ("learn/cox-ph.md", "ascii/cox-ph.txt", None,
     "Cox PH — Survival Analysis",
     "Model the hazard of leaving over time, not just whether someone leaves."),
    ("learn/causal-uplift.md", "ascii/causal-uplift.txt", None,
     "Causal Uplift — T-Learner",
     "Estimate the effect of an intervention per employee, not just the population."),
    ("learn/rag.md", "ascii/rag-fallback.txt", "diagrams/rag-fallback.png",
     "Multilingual RAG — Resilient Retrieval",
     "Hybrid index: OpenAI when reachable, hashing-based local fallback otherwise."),
    ("learn/fairness.md", "ascii/fairness.txt", "diagrams/fairness.png",
     "Fairness — Four-Fifths Rule",
     "Gate every prediction on the EEOC selection-rate ratio."),
    ("learn/nl-sql.md", "ascii/nl-sql.txt", None,
     "NL→SQL — Text-to-Database",
     "Hybrid pipeline: LLM-first, gold-question TF-IDF fallback."),
]


# --------------------------------------------------------------- low-level
def _set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_para_shading(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _add_horizontal_rule(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE \\* MERGEFORMAT"
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.color.rgb = MUTED
    run.font.size = Pt(9)


def _strip_inline_md(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"<([^>]+)>", r"\1", text)
    return text


def _add_inline_runs(paragraph, text, *, base_size=11, base_color=BODY, font="Calibri"):
    """Render **bold**, *italic*, `code` as proper runs."""
    text = _strip_inline_md(text)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = font
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            r.font.color.rgb = INK
            r.font.name = font
            r.font.size = Pt(base_size)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = CODE_INK
        else:  # *italic*
            r = paragraph.add_run(token[1:-1])
            r.italic = True
            r.font.color.rgb = base_color
            r.font.name = font
            r.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = font
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color


# --------------------------------------------------------------- builders
def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: INK, 2: LAVENDER, 3: PEACH}
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = colors.get(level, INK)
    if level == 1:
        _add_horizontal_rule(p)
    return p


def add_body(doc, text, *, italic=False, color=BODY, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if italic:
        r = p.add_run(_strip_inline_md(text))
        r.italic = True
        r.font.color.rgb = color
        r.font.size = Pt(size)
        r.font.name = "Calibri"
    else:
        _add_inline_runs(p, text, base_size=size, base_color=color)
    return p


def add_bullet(doc, text, *, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    _add_inline_runs(p, text)
    return p


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _set_para_shading(p, TINT_HEX)
    r = p.add_run("“ ")
    r.font.size = Pt(13)
    r.font.color.rgb = LAVENDER
    _add_inline_runs(p, text, base_color=BODY)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    _set_para_shading(p, CODE_BG_HEX)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = CODE_INK
    return p


def add_md_table(doc, lines):
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in lines]
    rows = [r for r in rows if not all(set(c) <= set("- :") for c in r)]
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.autofit = True
    for i, r in enumerate(rows):
        for j, c in enumerate(r):
            cell = tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(_strip_inline_md(c))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                run.font.color.rgb = INK
                _set_cell_shading(cell, TINT_HEX)
            else:
                run.font.color.rgb = BODY


def add_image(doc, path, *, caption=None, width_in=6.2):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED


def render_markdown(doc, md_text):
    """Render a markdown blob into the document."""
    # Strip YAML frontmatter
    if md_text.startswith("---"):
        end = md_text.find("\n---", 3)
        if end > 0:
            md_text = md_text[end + 4:].lstrip("\n")

    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                add_code(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        # tables
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.search(r"\|\s*-", lines[i + 1]):
            buf = [line]
            j = i + 1
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                buf.append(lines[j])
                j += 1
            add_md_table(doc, buf)
            i = j
            continue
        # block math $$ ... $$
        if line.strip().startswith("$$"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().endswith("$$"):
                buf.append(lines[i])
                i += 1
            add_code(doc, "\n".join(buf))
            i += 1
            continue
        # headings
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("> "):
            add_blockquote(doc, line[2:].strip())
        elif line.lstrip().startswith(("- ", "* ", "+ ")):
            add_bullet(doc, line.lstrip()[2:].strip())
        elif re.match(r"^\s*\d+\.\s+", line):
            m = re.match(r"^\s*\d+\.\s+(.*)$", line)
            if m:
                add_bullet(doc, m.group(1), ordered=True)
        elif re.match(r"^---+\s*$", line):
            doc.add_paragraph()  # spacer
        elif line.strip() == "":
            pass
        else:
            add_body(doc, line)
        i += 1


# --------------------------------------------------------------- TOC
def add_word_toc(doc):
    """Insert a Word-native TOC field. The TOC will populate when the
    user opens the document in Word and clicks 'Update Field'."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here in Word and choose 'Update Field' to populate the table of contents."
    run._r.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.color.rgb = MUTED
    run.font.size = Pt(10)
    run.italic = True


# --------------------------------------------------------------- main
def main():
    doc = Document()

    # Page setup
    s = doc.sections[0]
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    s.top_margin = Inches(0.85)
    s.bottom_margin = Inches(0.85)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)
    _add_page_number_footer(s)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY

    # ---------------- Cover page
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(120)
    r = cover.add_run("AttriSense")
    r.bold = True
    r.font.size = Pt(56)
    r.font.color.rgb = INK
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("The Beginner's Guide")
    r.font.size = Pt(28)
    r.font.color.rgb = LAVENDER

    rule = doc.add_paragraph()
    _add_horizontal_rule(rule)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_before = Pt(8)
    r = tag.add_run("Plain-language explainers for every technique in the dashboard.")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(60)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Version 2  ·  {date.today().isoformat()}")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    cover_banner = REPO / "docs" / "images" / "banner.png"
    if cover_banner.exists():
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(60)
        add_image(doc, cover_banner, width_in=6.5)

    doc.add_page_break()

    # ---------------- Foreword
    add_heading(doc, "Foreword", 1)
    add_body(doc,
             "AttriSense is a workforce-retention dashboard built around five model "
             "modalities — RandomForest with SMOTE, Cox proportional-hazards survival, "
             "EconML T-Learner causal uplift, multilingual RAG with provider fallback, "
             "and natural-language SQL with a TF-IDF safety net.")
    add_body(doc,
             "This guide is the offline companion to the live docs at docs/learn/. "
             "Each chapter starts with the intuition, sketches the math lightly, points "
             "to the file in production/src/attrisense/ that implements it, calls out "
             "the gotchas that bit us in production, and ends with a one-liner you can paste.")
    add_blockquote(doc,
                   "If you can read one paragraph, you can read AttriSense. "
                   "If you can read one chapter, you can ship it.")
    doc.add_page_break()

    # ---------------- TOC
    add_heading(doc, "Contents", 1)
    add_body(doc,
             "An interactive table of contents follows. Open this document in Word and "
             "press F9 (or right-click → Update Field) to populate the page numbers. "
             "A static index of chapters is also printed below.",
             italic=True, color=MUTED)
    add_word_toc(doc)
    add_body(doc, "")
    for idx, (_, _, _, title, key) in enumerate(SECTIONS, start=1):
        p = doc.add_paragraph()
        _add_inline_runs(p, f"**{idx}.  {title}** — {key}", base_color=BODY)
    doc.add_page_break()

    # ---------------- Chapters
    for idx, (md_rel, ascii_rel, diag_rel, title, key) in enumerate(SECTIONS, start=1):
        # Chapter title page
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(40)
        r = sub.add_run(f"Chapter {idx}")
        r.font.size = Pt(13)
        r.font.color.rgb = MUTED
        r.bold = True

        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(4)
        r = head.add_run(title)
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = INK
        _add_horizontal_rule(head)

        keyp = doc.add_paragraph()
        keyp.paragraph_format.space_before = Pt(8)
        r = keyp.add_run("Key idea  ·  ")
        r.bold = True
        r.font.color.rgb = LAVENDER
        r.font.size = Pt(12)
        r2 = keyp.add_run(key)
        r2.italic = True
        r2.font.color.rgb = BODY
        r2.font.size = Pt(12)

        # Render markdown body
        md_full = DOCS / md_rel
        if md_full.exists():
            doc.add_paragraph()
            render_markdown(doc, md_full.read_text())
        else:
            add_body(doc, f"(source missing: {md_rel})", italic=True, color=MUTED)

        # ASCII companion
        if ascii_rel:
            ascii_full = DOCS / ascii_rel
            if ascii_full.exists():
                add_heading(doc, "ASCII companion diagram", 3)
                add_code(doc, ascii_full.read_text())

        # Pixel pastel diagram PNG
        if diag_rel:
            png = DOCS / "images" / diag_rel.replace("diagrams/", "diagrams/")
            if png.exists():
                add_heading(doc, "Visual diagram", 3)
                add_image(doc, png, caption=f"Figure {idx}.  {title} — Pixel pastel diagram",
                          width_in=6.4)

        doc.add_page_break()

    # ---------------- Closing
    add_heading(doc, "Where to go next", 1)
    add_bullet(doc, "Read the full docs site under `docs/` (or `mkdocs serve`).")
    add_bullet(doc, "Read the production source under `production/src/attrisense/`.")
    add_bullet(doc, "Run the dashboard:  `streamlit run production/streamlit_app.py`.")
    add_bullet(doc, "Open the model card and fairness policy under `docs/ethics/`.")
    add_bullet(doc, "Read the IEEE conference paper at `paper/attrisense_ieee.pdf`.")

    add_heading(doc, "Repository layout", 2)
    add_md_table(doc, [
        "| Path | Purpose |",
        "|------|---------|",
        "| `production/streamlit_app.py` | The dashboard entry point |",
        "| `production/src/attrisense/` | Library: theme, identity, fairness, RAG, NL2SQL |",
        "| `scripts/` | Diagram generator, screenshot capture, beginner-guide builder |",
        "| `paper/` | IEEE conference paper (LaTeX, PDF, DOCX) |",
        "| `docs/` | mkdocs site sources, learning chapters, ASCII diagrams |",
        "| `archive/legacy/` | Older versions kept for reference |",
    ])

    add_heading(doc, "Credits", 2)
    add_body(doc,
             "Built by Sharada Dogiparthi. Pixel pastel theme, Aurora-coherent diagrams, "
             "and 5-modality model substrate. Pseudonymous Review_ID protects reviewers from "
             "identification bias.",
             italic=True, color=MUTED)

    doc.save(OUT)
    kb = OUT.stat().st_size // 1024
    pages = max(1, kb // 35)
    print(f"Wrote {OUT.relative_to(REPO)} ({kb} KB, ~{pages} pages estimated)")


if __name__ == "__main__":
    main()
