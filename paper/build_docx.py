# ---------------------------------------------------------------------------
# AttriSense — paper/build_docx.py
# ---------------------------------------------------------------------------
# Author : Sharada Dogiparthi <dogiparthi.sharada@gmail.com>
# Version: 1.0.0
# Date   : 2026-05-07
# License: MIT — see LICENSE in repo root.
# Copyright (c) 2026 Sharada Dogiparthi. All rights reserved.
# ---------------------------------------------------------------------------
"""Build attrisense_ieee.docx from attrisense_ieee.md (python-docx).

IEEE-conference style:
  - Title + author block: single column, centered.
  - Abstract + body: 2 columns, justified, Times New Roman 10pt.
  - Section headings: roman-numbered, all caps, centered, 10pt bold.
  - Subsection headings: italic, left, 10pt bold.
  - Tables, bullets, code blocks, and equations are rendered.

Source of truth = attrisense_ieee.md (kept in sync from the .tex via pandoc).
Run:
    .venv/bin/python build_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

HERE = Path(__file__).parent
MD = HERE / "attrisense_ieee.md"
OUT = HERE / "attrisense_ieee.docx"

ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI",
         "XII", "XIII", "XIV", "XV", "XVI"]


# --------------------------------------------------------------- helpers
def add_two_column_section(doc):
    s = doc.add_section(WD_SECTION.CONTINUOUS)
    cols = s._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")
    return s


def style_normal(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"<([^>]+)>", r"\1", text)
    text = text.replace("\\&", "&").replace("\\_", "_").replace("\\%", "%")
    return text


def _add_runs(paragraph, text, *, size=10, bold=False, italic=False, font="Times New Roman"):
    """Render text honouring **bold** and *italic* markers inline."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = font
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
        token = m.group(0)
        r = paragraph.add_run(token.strip("*`"))
        r.font.name = font
        r.font.size = Pt(size)
        if token.startswith("**"):
            r.bold = True
        elif token.startswith("*"):
            r.italic = True
        else:
            r.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic


def add_section_heading(doc, text, *, roman=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    label = f"{roman}. {text.upper()}" if roman else text.upper()
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"


def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"


def add_body(doc, text, *, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, _strip_md_inline(text))
    return p


def add_bullet(doc, text, *, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    _add_runs(p, _strip_md_inline(text))
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)


def add_md_table(doc, lines):
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in lines]
    rows = [r for r in rows if not all(set(c) <= set("- :") for c in r)]
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    for i, r in enumerate(rows):
        for j, c in enumerate(r):
            cell = tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(_strip_md_inline(c))
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True


# --------------------------------------------------------------- parser
def parse_md(md_text: str):
    """Split into (kind, payload) tokens.

    kinds: 'title', 'authors', 'abstract', 'keywords',
           'section', 'subsection', 'para', 'bullet', 'ordered',
           'code', 'table', 'rule'
    """
    # Strip the YAML frontmatter; we extract title + author manually.
    front = {}
    body = md_text
    if md_text.startswith("---"):
        end = md_text.find("\n---", 3)
        if end > 0:
            front_raw = md_text[3:end]
            body = md_text[end + 4:].lstrip("\n")
            cur_key = None
            for line in front_raw.splitlines():
                m = re.match(r"^(\w[\w-]*):\s*(.*)$", line)
                if m:
                    cur_key = m.group(1)
                    val = m.group(2).strip().strip('"')
                    front[cur_key] = val if val else []
                elif line.strip().startswith("- ") and cur_key:
                    if not isinstance(front[cur_key], list):
                        front[cur_key] = []
                    front[cur_key].append(line.strip()[2:].strip().strip('"'))

    tokens = []
    if front.get("title"):
        tokens.append(("title", front["title"]))
    if front.get("author"):
        authors = front["author"]
        if isinstance(authors, str):
            authors = [authors]
        tokens.append(("authors", authors))

    lines = body.splitlines()
    i = 0
    in_code = False
    code_buf = []
    para_buf = []

    def flush_para():
        if para_buf:
            tokens.append(("para", " ".join(para_buf).strip()))
            para_buf.clear()

    in_abstract = False
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("```"):
            if in_code:
                tokens.append(("code", "\n".join(code_buf)))
                code_buf = []
                in_code = False
            else:
                flush_para()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(ln)
            i += 1
            continue
        # md table
        if ln.lstrip().startswith("|") and i + 1 < len(lines) and re.search(r"\|\s*-", lines[i + 1]):
            flush_para()
            buf = [ln]
            j = i + 1
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                buf.append(lines[j])
                j += 1
            tokens.append(("table", buf))
            i = j
            continue
        # heading
        m = re.match(r"^(#+)\s+(.*?)\s*$", ln)
        if m:
            flush_para()
            level = len(m.group(1))
            heading = m.group(2).strip()
            heading_clean = re.sub(r"^\d+(\.\d+)*\s+", "", heading)
            if heading_clean.lower() == "abstract":
                in_abstract = True
                tokens.append(("abstract_marker", None))
            elif heading_clean.lower().startswith("keywords"):
                in_abstract = False
                tokens.append(("keywords_marker", None))
            elif level == 1 or (level == 2 and heading.startswith(tuple("0123456789"))):
                in_abstract = False
                tokens.append(("section", heading_clean))
            else:
                tokens.append(("subsection", heading_clean))
            i += 1
            continue
        # horizontal rule
        if re.match(r"^---+\s*$", ln):
            flush_para()
            i += 1
            continue
        # bullets
        m = re.match(r"^\s*[-*+]\s+(.*)$", ln)
        if m:
            flush_para()
            tokens.append(("bullet", m.group(1)))
            i += 1
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        if m:
            flush_para()
            tokens.append(("ordered", m.group(1)))
            i += 1
            continue
        # blank line ends paragraph
        if not ln.strip():
            flush_para()
            i += 1
            continue
        para_buf.append(ln.strip())
        i += 1
    flush_para()
    return front, tokens


# --------------------------------------------------------------- main
def main():
    md_text = MD.read_text(encoding="utf-8")
    front, tokens = parse_md(md_text)

    doc = Document()
    s0 = doc.sections[0]
    s0.page_width = Inches(8.5)
    s0.page_height = Inches(11)
    s0.top_margin = Inches(0.75)
    s0.bottom_margin = Inches(1.0)
    s0.left_margin = Inches(0.625)
    s0.right_margin = Inches(0.625)
    style_normal(doc)

    # --- Title block (single column)
    title_text = front.get("title") or "AttriSense"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_strip_md_inline(title_text))
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(12)

    for author in front.get("author", []):
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = ap.add_run(author)
        ar.bold = True
        ar.font.size = Pt(11)
        ar.font.name = "Times New Roman"
    if front.get("author"):
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- 2-column section
    add_two_column_section(doc)

    section_idx = 0
    in_abstract = False
    keywords_pending = False

    for kind, payload in tokens:
        if kind == "abstract_marker":
            in_abstract = True
            continue
        if kind == "keywords_marker":
            in_abstract = False
            keywords_pending = True
            continue
        if kind == "section":
            section_idx += 1
            roman = ROMAN[section_idx - 1] if section_idx <= len(ROMAN) else str(section_idx)
            add_section_heading(doc, payload, roman=roman)
            continue
        if kind == "subsection":
            add_subsection_heading(doc, payload)
            continue
        if kind == "para":
            if in_abstract:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run("Abstract\u2014")
                r.bold = True
                r.italic = True
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"
                _add_runs(p, _strip_md_inline(payload), size=9)
                in_abstract = False
            elif keywords_pending:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run("Index Terms\u2014")
                r.bold = True
                r.italic = True
                r.font.size = Pt(9)
                _add_runs(p, _strip_md_inline(payload), size=9, italic=False)
                keywords_pending = False
            else:
                add_body(doc, payload)
            continue
        if kind == "bullet":
            add_bullet(doc, payload)
            continue
        if kind == "ordered":
            add_bullet(doc, payload, ordered=True)
            continue
        if kind == "code":
            add_code(doc, payload)
            continue
        if kind == "table":
            add_md_table(doc, payload)
            continue

    doc.save(OUT)
    pages_hint = max(1, OUT.stat().st_size // 35000)
    print(f"Wrote {OUT.relative_to(HERE.parent)} "
          f"({OUT.stat().st_size // 1024} KB, ~{pages_hint} pages estimated)")


if __name__ == "__main__":
    main()
